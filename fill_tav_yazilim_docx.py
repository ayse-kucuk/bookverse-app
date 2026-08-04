# -*- coding: utf-8 -*-
"""
tav_staj_defteri_sozel.md icerigini bos sablondan uretilen
tav_yazilim_staj.docx dosyasina gun gun (Gun 1..Gun 20) yerlestirir.
"""
import re
import copy
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

MD_PATH = "tav_staj_defteri_sozel.md"
DOCX_IN = "yazilim_staj_yedek.docx"
DOCX_OUT = "tav_yazilim_staj.docx"


def business_days(start, end, holidays):
    days = []
    d = start
    while d <= end:
        if d.weekday() < 5 and d not in holidays:
            days.append(d)
        d += datetime.timedelta(days=1)
    return days


# Yazilim staji: 29 Haziran 2026 - 27 Temmuz 2026, hafta ici, 15 Temmuz haric
DATES = business_days(
    datetime.date(2026, 6, 29),
    datetime.date(2026, 7, 27),
    {datetime.date(2026, 7, 15)},
)
assert len(DATES) == 20, f"Beklenen 20 gun, bulunan {len(DATES)}"


def parse_markdown(path):
    with open(path, encoding="utf-8") as f:
        text = f.read()

    day_pattern = re.compile(r"^### Gün (\d+): (.+)$", re.MULTILINE)
    matches = list(day_pattern.finditer(text))
    days = []
    for idx, m in enumerate(matches):
        day_no = int(m.group(1))
        title = m.group(2).strip()
        body_start = m.end()
        body_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        body = text[body_start:body_end]
        body = body.split("\n---")[0].strip()

        blocks = [b.strip() for b in re.split(r"\n\s*\n", body) if b.strip()]

        teorik = ""
        teknik_p1 = ""
        code = ""
        teknik_p2_parts = []
        zorluk = ""

        stage = None
        for b in blocks:
            if b.startswith("- **Teorik"):
                teorik = re.sub(r"^- \*\*.+?:\*\*\s*", "", b)
                stage = "teorik"
            elif b.startswith("- **Teknik"):
                teknik_p1 = re.sub(r"^- \*\*.+?:\*\*\s*", "", b)
                stage = "teknik"
            elif b.startswith("```"):
                code_lines = b.split("\n")[1:-1]
                code = "\n".join(code_lines)
                stage = "code"
            elif b.startswith("- **Karşılaşılan"):
                zorluk = re.sub(r"^- \*\*.+?:\*\*\s*", "", b)
                stage = "zorluk"
            else:
                if stage in ("teknik", "code"):
                    teknik_p2_parts.append(b)

        days.append(
            {
                "day": day_no,
                "title": title,
                "teorik": teorik,
                "teknik_p1": teknik_p1,
                "code": code,
                "teknik_p2": " ".join(teknik_p2_parts),
                "zorluk": zorluk,
            }
        )
    return days


DAYS = parse_markdown(MD_PATH)
assert len(DAYS) == 20, f"Beklenen 20 gun parse edildi, bulunan {len(DAYS)}"


def add_rich_text(paragraph, text, base_bold=False, size=12, font_name=None):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(size)
            run.bold = base_bold
        else:
            run = paragraph.add_run(part)
            if font_name:
                run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = base_bold


def new_blank_paragraph(template_p_element, doc):
    new_el = copy.deepcopy(template_p_element)
    for child in list(new_el):
        if child.tag != qn("w:pPr"):
            new_el.remove(child)
    from docx.text.paragraph import Paragraph

    return Paragraph(new_el, doc)


def make_section_paragraph(template_el, doc, label, text, justify=True):
    p = new_blank_paragraph(template_el, doc)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(12)
    add_rich_text(p, text, base_bold=False, size=12)
    return p


def make_plain_paragraph(template_el, doc, text, justify=True):
    p = new_blank_paragraph(template_el, doc)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_rich_text(p, text, base_bold=False, size=12)
    return p


def make_code_paragraph(template_el, doc, code_text):
    p = new_blank_paragraph(template_el, doc)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = code_text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else "\u00A0")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        if i < len(lines) - 1:
            run.add_break()
    return p


def make_pagebreak_paragraph(template_el, doc):
    p = new_blank_paragraph(template_el, doc)
    run = p.add_run("")
    run.add_break(WD_BREAK.PAGE)
    return p


def set_date_in_header_table(table, date_str):
    date_cell = table.rows[0].cells[1]
    para = date_cell.paragraphs[0]
    last_run = para.runs[-1]
    last_run.text = f" {date_str}"


def set_title_in_header_table(table, doc, title):
    konu_cell = table.rows[0].cells[0]
    para = konu_cell.paragraphs[0]
    run = para.add_run(title)
    run.bold = False
    run.font.size = Pt(12)


def set_paragraph_text(paragraph, text, size=12):
    if not paragraph.runs:
        r = paragraph.add_run(text)
        r.font.size = Pt(size)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def fill_cover_tables(doc, start_str, end_str, total_days):
    ADI_SOYADI = "Ayşe Küçük"
    OGRENCI_NO = "2311012075"
    YARIYIL = "2025 / 2026"
    KURUM = "TAV Bilişim / TAV Technologies"
    STAJ_KONUSU = "Yazılım Geliştirme ve Web Teknolojileri"

    t0 = doc.tables[0]
    set_paragraph_text(t0.rows[3].cells[1].paragraphs[0], ADI_SOYADI)
    set_paragraph_text(t0.rows[4].cells[1].paragraphs[0], OGRENCI_NO)
    set_paragraph_text(t0.rows[5].cells[1].paragraphs[0], STAJ_KONUSU)

    t2 = doc.tables[2]
    set_paragraph_text(t2.rows[1].cells[1].paragraphs[0], ADI_SOYADI)
    set_paragraph_text(t2.rows[2].cells[1].paragraphs[0], OGRENCI_NO)
    set_paragraph_text(t2.rows[3].cells[1].paragraphs[0], YARIYIL)
    set_paragraph_text(t2.rows[4].cells[1].paragraphs[0], STAJ_KONUSU)

    kurum_cell = t2.rows[5].cells[1]
    set_paragraph_text(kurum_cell.paragraphs[1], KURUM)

    start_parts = start_str.split("/")
    end_parts = end_str.split("/")
    set_paragraph_text(
        t2.rows[6].cells[1].paragraphs[0],
        f"{start_parts[0]} /{start_parts[1]}/{start_parts[2]}",
    )
    set_paragraph_text(
        t2.rows[7].cells[1].paragraphs[0],
        f"{end_parts[0]} /{end_parts[1]}/ {end_parts[2]}",
    )
    set_paragraph_text(t2.rows[8].cells[1].paragraphs[0], f"{total_days} İş Günü")


def fix_section_break_to_nextpage(doc):
    body = doc.element.body
    final_sectPr = body.find(qn("w:sectPr"))
    type_el = final_sectPr.find(qn("w:type"))
    if type_el is not None:
        type_el.set(qn("w:val"), "nextPage")
    else:
        new_type = final_sectPr.makeelement(qn("w:type"), {qn("w:val"): "nextPage"})
        final_sectPr.insert(0, new_type)

    pgNumType = final_sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        pgNumType.set(qn("w:start"), "1")


def fill_summary_table(doc, day_info_list):
    t3 = doc.tables[3]
    for i, info in enumerate(day_info_list):
        row = t3.rows[i + 1]
        date_cell, konu_cell, sayfa_cell = row.cells[0], row.cells[1], row.cells[2]
        date_cell.paragraphs[0].runs[0].text = info["date_str"]
        konu_para = konu_cell.paragraphs[0]
        run = konu_para.add_run(info["title"])
        run.font.size = Pt(11)
        sayfa_para = sayfa_cell.paragraphs[0]
        run2 = sayfa_para.add_run(str(i + 1))
        run2.font.size = Pt(11)

    total_row = t3.rows[27]
    total_cell = total_row.cells[0]
    total_cell.paragraphs[0].runs[0].text = "20"


def main():
    doc = Document(DOCX_IN)
    body = doc.element.body
    day_info_for_summary = []

    for i, day in enumerate(DAYS):
        table = doc.tables[4 + i]
        date_str = DATES[i].strftime("%d/%m/%Y")

        set_date_in_header_table(table, date_str)
        set_title_in_header_table(table, doc, day["title"])
        day_info_for_summary.append({"date_str": date_str, "title": day["title"]})

        children = list(body.iterchildren())
        table_idx = children.index(table._tbl)
        is_last_real_day = i == len(DAYS) - 1
        physical_next_table = doc.tables[4 + i + 1] if (4 + i + 1) < len(doc.tables) else None
        if physical_next_table is not None:
            next_idx = children.index(physical_next_table._tbl)
        else:
            next_idx = len(children)

        gap_elements = children[table_idx + 1 : next_idx]
        if not gap_elements:
            raise RuntimeError(f"Gun {i+1} icin bos paragraf alani bulunamadi")
        template_el = gap_elements[0]

        new_paragraphs = []
        new_paragraphs.append(
            make_section_paragraph(template_el, doc, "Teorik Bilgi ve Amaç", day["teorik"])
        )
        new_paragraphs.append(new_blank_paragraph(template_el, doc))
        new_paragraphs.append(
            make_section_paragraph(
                template_el, doc, "Teknik Uygulama ve Mimari Kararlar", day["teknik_p1"]
            )
        )
        if day["code"]:
            new_paragraphs.append(make_code_paragraph(template_el, doc, day["code"]))
        if day["teknik_p2"]:
            new_paragraphs.append(make_plain_paragraph(template_el, doc, day["teknik_p2"]))
        new_paragraphs.append(new_blank_paragraph(template_el, doc))
        new_paragraphs.append(
            make_section_paragraph(
                template_el, doc, "Karşılaşılan Zorluk ve Çözüm", day["zorluk"]
            )
        )
        if not is_last_real_day:
            new_paragraphs.append(make_pagebreak_paragraph(template_el, doc))

        for el in gap_elements:
            el.getparent().remove(el)

        ref = table._tbl
        for p in new_paragraphs:
            ref.addnext(p._p)
            ref = p._p

    if len(doc.tables) > 4 + len(DAYS):
        first_extra_table = doc.tables[4 + len(DAYS)]
        children = list(body.iterchildren())
        extra_start_idx = children.index(first_extra_table._tbl)
        for el in children[extra_start_idx:-1]:
            el.getparent().remove(el)

    fix_section_break_to_nextpage(doc)
    fill_summary_table(doc, day_info_for_summary)
    fill_cover_tables(
        doc,
        DATES[0].strftime("%d/%m/%Y"),
        DATES[-1].strftime("%d/%m/%Y"),
        len(DATES),
    )

    doc.save(DOCX_OUT)
    print(f"Basarili: {len(DAYS)} gun '{DOCX_OUT}' dosyasina yazildi.")
    for info in day_info_for_summary:
        print(f"  {info['date_str']} — {info['title'][:60]}")


if __name__ == "__main__":
    main()
